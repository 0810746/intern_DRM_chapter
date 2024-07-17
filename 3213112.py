from selenium import webdriver
from selenium.webdriver.common.by import By
import pandas as pd

# 初始化 WebDriver
driver = webdriver.Chrome()  # 確保你已經安裝了 ChromeDriver

# 打開目標網頁
url = '你的目標網址'
driver.get(url)

# 找到所有的 <tr> 元素
rows = driver.find_elements(By.TAG_NAME, 'tr')

# 初始化列表來存儲表格數據
table_data = []

# 迭代每一行
for row in rows:
    cells = row.find_elements(By.XPATH, ".//td|.//th")
    row_data = []
    for cell in cells:
        cell_text = cell.text.strip()
        rowspan = int(cell.get_attribute('rowspan') or 1)
        colspan = int(cell.get_attribute('colspan') or 1)
        for _ in range(rowspan):
            for _ in range(colspan):
                row_data.append(cell_text)
    table_data.append(row_data)


import re
from docx import Document

# 打開 Word 文件
input_path = r'C:\Users\steve\Desktop\intern_practice\small\in_data\論文.docx'
document = Document(input_path)

# 定義匹配模式
pattern = re.compile(r'\bkjh\.\d+\.\d+\.\d+\.kjh\.\d+\.kh\b')

# 遍歷所有表格
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            # 檢查單元格內容是否匹配模式
            if pattern.match(cell.text):
                print(f'匹配內容: {cell.text} 在第 {document.tables.index(table) + 1} 張表, 第 {table.rows.index(row) + 1} 行, 第 {row.cells.index(cell) + 1} 列')


# 將數據轉換為 DataFrame
df = pd.DataFrame(table_data)

# 迭代處理合併的儲存格
for i, row in df.iterrows():
    for j, cell in enumerate(row):
        if cell != "":
            colspan = int(rows[i].find_elements(By.XPATH, ".//td|.//th")[j].get_attribute('colspan') or 1)
            rowspan = int(rows[i].find_elements(By.XPATH, ".//td|.//th")[j].get_attribute('rowspan') or 1)
            for k in range(1, colspan):
                df.iat[i, j + k] = cell
            for l in range(1, rowspan):
                df.iat[i + l, j] = cell

# 將 DataFrame 保存到 Excel 文件
output_file = 'output.xlsx'
df.to_excel(output_file, index=False, header=False)

print(f"Successfully saved table data to {output_file}")


#====================================================================
# 找到 <p> 元素中的所有 <span> 和 <br> 元素
p_element = driver.find_element(By.XPATH, '//p')
span_elements = p_element.find_elements(By.XPATH, './span | ./br')

# 遍歷元素並組合文本，保留 <br> 標籤
combined_text = ""
for element in span_elements:
    if element.tag_name == 'span':
        combined_text += element.text
    elif element.tag_name == 'br':
        combined_text += '<br>'


# 找到 <p> 元素
p_element = driver.find_element(By.XPATH, '//p')

# 遍歷 <p> 元素中的所有子元素
combined_text = ""
for child in p_element.find_elements(By.XPATH, './*'):
    if child.tag_name == 'span':
        combined_text += child.text
    elif child.tag_name == 'br':
        combined_text += '\n'  # 或者使用 '<br>' 來保留 HTML 標籤
