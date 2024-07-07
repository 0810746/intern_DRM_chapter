from docx import Document
from docx.oxml import OxmlElement
from copy import deepcopy


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None

input_path = r'C:\Users\steve\Desktop\intern_practice\DRMDB\data\only_table.docx'#r'C:\Users\steve\Desktop\intern_practice\DRMDB\data\論文.docx'
document = Document(input_path)
output_path = r'C:\Users\steve\Desktop\intern_practice\DRMDB\data\table.docx'
#刪除第一頁
for shape in document.element.xpath('//w:sdt'):
    shape.getparent().remove(shape)
    break

document.save(output_path)

document = Document(output_path)
save_tables = []
temp = 0
for table in document.tables:
    target_header = ['構成要件\n罪名', '身分', '手段要件', '結果要件', '主觀要件']
    #判斷式
    first_row = [cell.text.strip() for cell in table.rows[0].cells]
    if first_row == target_header:
        save_tables.append(temp)
    temp +=1

print(save_tables)
document = Document(output_path)
temp = 0
for i  in range(len(document.tables)):
    if i not in save_tables:
        table = document.tables[i - temp]
        table._element.getparent().remove(table._element)
        temp += 1   
document.save(output_path)

document = Document(output_path)
for p in document.paragraphs:
    delete_paragraph(p)

tables = document.tables
if len(tables) > 1:
    for table in tables[:-1]:
        new_paragraph = OxmlElement("w:p")
        table._element.addnext(new_paragraph)

document.save(output_path)

#存檔 html-----------------------------------------------------------------------
# from docx import Document
# from docx2html import convert

# def convert_docx_to_html(input_path, output_html):
#     document = Document(input_path)
#     # 将文档转换为 HTML
#     html_content = convert(input_path)

#     # 保存 HTML
#     with open(output_html, 'w', encoding='utf-8') as f:
#         f.write(html_content)

# convert_docx_to_html(output_path, output_html)

# import mammoth

# input_path = 'path_to_your_word_document.docx'
# output_html_path = 'output_file.html'

# with open(input_path, "rb") as docx_file:
#     result = mammoth.convert_to_html(docx_file)
#     html_content = result.value

# with open(output_html_path, "w", encoding="utf-8") as html_file:
#     html_file.write(html_content)

# print(f"Successfully converted {input_path} to {output_html_path}")





# #存excel pandas-------------------------------------------------------
# # 将表格数据转换为 DataFrame
# dfs = []
# for table in save_tables:
#     data = []
#     for row in table.rows:
#         row_data = [cell.text.strip() for cell in row.cells]
#         data.append(row_data)
#     df = pd.DataFrame(data[1:], columns=data[0])  # 假设第一行是表头
#     # df.set_index(data[0][0], inplace=True)  # 将第一列设置为索引
#     dfs.append(df)

# # 将每个 DataFrame 存储到 Excel 文件中
# output_excel = r'C:\Users\steve\Desktop\intern_practice\DRMDB\data\output_tables.xlsx'
# with pd.ExcelWriter(output_excel) as writer:
#     for i, df in enumerate(dfs):
#         sheet_name = f'Table_{i+1}'  # 设定每个表格的 sheet 名称
#         df.to_excel(writer, sheet_name=sheet_name, index=False)


