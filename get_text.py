from docx import Document
import pandas as pd
from copy import deepcopy

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None

input_path = r'C:\Users\steve\Desktop\intern_practice\DRMDB\data\論文.docx'
document = Document(input_path)
output_path = r'C:\Users\steve\Desktop\intern_practice\DRMDB\data\text.txt'
#刪除第一頁
for shape in document.element.xpath('//w:sdt'):
    shape.getparent().remove(shape)
    break
document.save(output_path)
#刪除表格
for i in range(len(document.tables)):
    document = Document(output_path)
    table = document.tables[0]
    table._element.getparent().remove(table._element)
    document.save(output_path)


#%% Function to extract text from Word document--------------------------------------
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)
    

# Extract text from the Word document
try:
    extracted_text = extract_text_from_docx(input_path)
    # Write extracted text to a text file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(extracted_text)
    print(f"Successfully converted {input_path} to {output_path}")
except Exception as e:
    print(f"Error converting {input_path} to text: {e}")


# %%
