import pandas as pd
from docx import Document

#%% 存檔 excel--------------------------------------------------------------------------
input_path = r'C:\Users\steve\Desktop\intern_practice\DRMDB\data\table.docx'
document = Document(input_path)
output_path = r'C:\Users\steve\Desktop\intern_practice\DRMDB\data\output_tables.xlsx'

save_tables = []
for table in document.tables:
    target_header = ['構成要件\n罪名', '身分', '手段要件', '結果要件', '主觀要件']
    #判斷式
    first_row = [cell.text.strip() for cell in table.rows[0].cells]
    if first_row == target_header:
        save_tables.append(table)

#%%
dfs = []
for table in save_tables:
    data = []
    # 保存第一行的样式
    first_row_format = [cell.text.strip() for cell in table.rows[0].cells]
    for i, row in enumerate(table.rows):
        row_data = [cell.text.strip() for cell in row.cells]
        if i == 0:
            row_data_formatted = first_row_format  # 使用保存的第一行样式
        else:
            row_data_formatted = row_data
        data.append(row_data_formatted)
    
    df = pd.DataFrame(data[1:], columns=data[0])  # 假设第一行是表头
    dfs.append(df)

# 将每个 DataFrame 存储到 Excel 文件中
with pd.ExcelWriter(output_path) as writer:
    for i, df in enumerate(dfs):
        sheet_name = f'Table_{i+1}'  # 设定每个表格的 sheet 名称
        df.to_excel(writer, sheet_name=sheet_name, index=False)

print(f"Successfully converted {input_path} to {output_path}")

#%% 存excel pandas-------------------------------------------------------
# 将表格数据转换为 DataFrame
dfs = []
for table in save_tables:
    data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        data.append(row_data)
    df = pd.DataFrame(data[1:], columns=data[0])  # 假设第一行是表头
    # df.set_index(data[0][0], inplace=True)  # 将第一列设置为索引
    dfs.append(df)

# 将每个 DataFrame 存储到 Excel 文件中
output_excel = r'C:\Users\steve\Desktop\intern_practice\DRMDB\data\output_tables.xlsx'
with pd.ExcelWriter(output_excel) as writer:
    for i, df in enumerate(dfs):
        sheet_name = f'Table_{i+1}'  # 设定每个表格的 sheet 名称
        df.to_excel(writer, sheet_name=sheet_name, index=False)

# %% jason
import json
from docx import Document

# 输入和输出文件路径
input_path = r'C:\Users\steve\Desktop\intern_practice\DRMDB\data\table.docx'
output_path = r'C:\Users\steve\Desktop\intern_practice\DRMDB\data\output_tables.json'

# 读取 Word 文档
document = Document(input_path)

# 存储目标表格和目标表头
save_tables = []
target_header = ['構成要件\n罪名', '身分', '手段要件', '結果要件', '主觀要件']

# 遍历所有表格，找到符合条件的表格
for table in document.tables:
    first_row = [cell.text.strip() for cell in table.rows[0].cells]
    if first_row == target_header:
        save_tables.append(table)

# 将符合条件的表格数据转换为 JSON 格式
tables_data = []
for table in save_tables:
    table_data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)
    tables_data.append(table_data)

# 写入 JSON 文件
with open(output_path, 'w', encoding='utf-8') as f:
    json.dump(tables_data, f, ensure_ascii=False, indent=4)

print(f"Successfully converted {input_path} to {output_path}")

# %%
import pandas as pd
from docx import Document

# Input and output file paths
input_path = r'C:\Users\steve\Desktop\intern_practice\DRMDB\data\table.docx'
output_folder = r'C:\Users\steve\Desktop\intern_practice\DRMDB\data'

# Read the Word document
document = Document(input_path)

# Define target header for identifying tables
target_header = ['構成要件\n罪名', '身分', '手段要件', '結果要件', '主觀要件']

# Function to extract table data from a Word table
def extract_table_data(table):
    table_data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)
    return table_data

# List to store all table dataframes
table_dfs = []

# Process each table in the document
for table in document.tables:
    first_row = [cell.text.strip() for cell in table.rows[0].cells]
    if first_row == target_header:
        table_data = extract_table_data(table)
        df = pd.DataFrame(table_data[1:], columns=table_data[0])
        table_dfs.append(df)

# Write each DataFrame to a CSV file
for idx, df in enumerate(table_dfs):
    output_csv_path = f"{output_folder}\\table_{idx + 1}.csv"
    df.to_csv(output_csv_path, index=False, encoding='utf-8-sig')  # utf-8-sig for Excel compatibility

    print(f"Converted table {idx + 1} to CSV: {output_csv_path}")

print("All tables converted to CSV successfully.")
