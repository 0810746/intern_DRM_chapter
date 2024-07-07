
#%%
from docx import Document
import pandas as pd
from copy import deepcopy

input_path = r'C:\Users\steve\Desktop\intern_practice\small\in_data\【論文】.docx'
document = Document(input_path)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None

def chapter_paragraph_location(document, df):
    temp = 0
    for i, parag in enumerate(document.paragraphs):
        if parag.text == df.loc[temp, 'name']:
            df.loc[temp, 'location'] = i
            temp += 1
        if temp == len(df)-1:
            break

def chapter_paragraph_location_end(document, df):
    for i, parag in enumerate(document.paragraphs):
        temp = i
    df.loc[len(df)-1, 'location'] = temp-1


location = [0, 0, 0, 0, 0, 0] 
chapter = ['第一章', '第二章', '第三章', '第四章', '第五章', 'end']
name = ['緒論','強迫勞動之定義、規範與現象探討', '歐盟與其會員國法國之強制性人權盡職調查立法之探討', '我國供應鏈企業責任立法之考察比較', '結論與建議', '文檔最後']
data = {'location' : location,'chapter' : chapter,'name' : name}
chlo = pd.DataFrame(data)

chapter_paragraph_location(document , chlo)
print(chlo.location)


chapter_number = 5
output_path = r'C:\Users\steve\Desktop\intern_practice\small\out_data\{}.docx'.format(chlo.chapter[chapter_number-1])

#刪除第一頁
for shape in document.element.xpath('//w:sdt'):
    shape.getparent().remove(shape)
    break



table = document.add_table(rows=2, cols=2)
for row in table.rows:
    for cell in row.cells:
        cell.text = 'separate'

separate = deepcopy(document.tables[len(document.tables)-1])
document.save(output_path)    

for tindex in range(len(chlo)-1):
    document = Document(output_path)
    for i,parag in enumerate(document.paragraphs):
        if i == chlo.location[tindex] - 1 + tindex:
            parag.insert_paragraph_before(' ')
            parag._p.addnext(separate._element)
            document.save(output_path)


chapter_paragraph_location(document , chlo)
chapter_paragraph_location_end(document , chlo)
print(chlo.location)


#%%
# chapter_number = 1
document = Document(output_path)
start = chlo.location[chapter_number - 1] - 1
end = chlo.location[chapter_number] - 1
print(start,end)


temp = 0
for p in document.paragraphs:
    if temp <= start or temp >= end :
        delete_paragraph(p)    
    temp+=1

document.save(output_path)
document = Document(output_path)

sep_list = []
for i in range(len(document.tables)):
    table = document.tables[i]
    # print(table.rows[0].cells[0].text)
    if table.rows[0].cells[0].text == "separate":
        sep_list.append(i)
print(sep_list)

temp = len(document.tables)
print(temp)


for i in range(temp  - sep_list[chapter_number]):
    document = Document(output_path)
    temp = len(document.tables)
    table = document.tables[temp-1]
    table._element.getparent().remove(table._element)
    document.save(output_path)

for i in range(sep_list[chapter_number-1] + 1):
    document = Document(output_path)
    table = document.tables[0]
    table._element.getparent().remove(table._element)
    document.save(output_path)




# output_path = r'C:\Users\steve\Desktop\intern_practice\small\out_data\{}.docx'.format(s.index[0])
document.save(output_path)
# %%
